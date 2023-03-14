import logging

from aiogram import types
from aiogram.dispatcher.filters.builtin import CommandStart
from aiogram.dispatcher import FSMContext
from keyboards.inline.main_inline import *
from keyboards.inline.menu_button import *
from loader import dp, bot
from utils.db_api.database import *
import datetime
from aiogram.types import ReplyKeyboardRemove
from geopy.geocoders import Nominatim
from aiogram.utils.deep_linking import decode_payload, get_start_link
import re
import requests
from docx import Document
from docx2pdf import convert
from get_valute import valyuta_kurslari
import os


def isValid(s):
    Pattern = re.compile("(0|91)?[7-9][0-9]{9}")
    return Pattern.match(s)


async def send_sms(otp, phone):
    username = 'bestbrok'
    password = 'tM4!-hmV52Z@'
    sms_data = {
        "messages": [{"recipient": f"{phone}", "message-id": "abc000000003", "sms": {"originator": "3700", "content":
        {"text": f"Ваш код подтверждения для BEST BROK BOT: {otp}"}}}]}
    url = "http://91.204.239.44/broker-api/send"
    res = requests.post(url=url, auth=(username, password), json=sms_data)
    print(res)


async def generateOTP():
    return random.randint(111111, 999999)


@dp.message_handler(commands=["menu"], state="*")
async def add_datas(message: types.Message, state: FSMContext):
    user = await get_user(message.from_id)
    if user is not None:
        if user.lang:
            lang = await get_lang(message.from_user.id)
            if user.phone:
                markup = await user_menu(lang)
                if lang == "uz":
                    await message.answer("Botimizga xush kelibsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
                elif lang == "ru":
                    await message.answer("Добро пожаловать в наш бот. Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
                elif lang == "en":
                    await message.answer("Welcome to our bot. Please select the desired section 👇", reply_markup=markup, protect_content=True)
                await state.set_state("get_category")
            else:
                markup = await back_to_keyboard(lang)
                if lang == "uz":
                    await message.answer("Iltimos ismingizni kiriting 👇", reply_markup=markup, protect_content=True)
                elif lang == "ru":
                    await message.answer("Пожалуйста, введите ваше имя 👇", reply_markup=markup, protect_content=True)
                elif lang == "en":
                    await message.answer("Please enter your name 👇", reply_markup=markup, protect_content=True)
                await state.set_state("get_name")
                
        else:
            markup =await language_keyboard()
            await message.answer(f"Assalomu alaykum, {message.from_user.first_name}👋. \nKerakli tilni tanlang 👇\n\nHello, {message.from_user.first_name}👋. \nChoose the language you need 👇\n\nЗдравствуйте, {message.from_user.first_name}👋. \nВыберите нужный язык 👇", 
                                reply_markup=markup, protect_content=True)
            await state.set_state("get_lang")


@dp.message_handler(commands=["add"], state="*")
async def add_datas(message: types.Message, state: FSMContext):
    add_data()


@dp.message_handler(lambda message: message.text in ["🏠 Asosiy menyu", "🏠 Main menu", "🏠 Главное меню"], state='*')
async def go_home(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    markup = await user_menu(lang)
    if lang == "uz":
        await message.answer("Botimizga xush kelibsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
    elif lang == "ru":
        await message.answer("Добро пожаловать в наш бот. Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
    elif lang == "en":
        await message.answer("Welcome to our bot. Please select the desired section 👇", reply_markup=markup, protect_content=True)
    await state.set_state("get_category")
 

@dp.message_handler(CommandStart(), state='*')
async def bot_start(message: types.Message, state: FSMContext):
    user = await get_user(message.from_id)
    if user is not None:
        if user.lang:
            lang = await get_lang(message.from_user.id)
            if user.phone:
                markup = await user_menu(lang)
                if lang == "uz":
                    await message.answer("Botimizga xush kelibsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
                elif lang == "ru":
                    await message.answer("Добро пожаловать в наш бот. Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
                elif lang == "en":
                    await message.answer("Welcome to our bot. Please select the desired section 👇", reply_markup=markup, protect_content=True)
                await state.set_state("get_category")
            else:
                markup = await back_to_keyboard(lang)
                if lang == "uz":
                    await message.answer("Iltimos ismingizni kiriting 👇", reply_markup=markup, protect_content=True)
                elif lang == "ru":
                    await message.answer("Пожалуйста, введите ваше имя 👇", reply_markup=markup, protect_content=True)
                elif lang == "en":
                    await message.answer("Please enter your name 👇", reply_markup=markup, protect_content=True)
                await state.set_state("get_name")
                
        else:
            markup =await language_keyboard()
            await message.answer(f"Assalomu alaykum, {message.from_user.first_name}👋. \nKerakli tilni tanlang 👇\n\nHello, {message.from_user.first_name}👋. \nChoose the language you need 👇\n\nЗдравствуйте, {message.from_user.first_name}👋. \nВыберите нужный язык 👇", 
                                reply_markup=markup, protect_content=True)
            await state.set_state("get_lang")
            
    else:
        args = message.get_args()
        payload = decode_payload(args)
        if payload != '':
            await add_user(user_id=message.from_user.id, referal_user=payload)
        else:
            await add_user(user_id=message.from_user.id, referal_user="no_referal")
        markup =await language_keyboard()
        await message.answer(f"Assalomu alaykum, {message.from_user.first_name}👋. \nKerakli tilni tanlang 👇\n\nHello, {message.from_user.first_name}👋. \nChoose the language you need 👇\n\nЗдравствуйте, {message.from_user.first_name}👋. \nВыберите нужный язык 👇", 
                            reply_markup=markup, protect_content=True)
        await state.set_state("get_lang")

@dp.message_handler(content_types=['document'], state='*')
async def bot_start(message: types.Message, state: FSMContext):
    await message.answer(text=message.document.file_id)


@dp.message_handler(state="get_lang")
async def get_language(message: types.Message, state: FSMContext):
    if message.text in ["🇺🇿 O'zbek tili", "🇺🇸 English", "🇷🇺 Русский язык"]:
        data = []
        if message.text == "🇺🇿 O'zbek tili":
            data = "uz"
        elif message.text == "🇺🇸 English":
            data = "en"
        elif message.text == "🇷🇺 Русский язык":
            data = "ru"
        user = await get_user(message.from_user.id)
        user.lang = data
        user.save()
        lang = await get_lang(message.from_user.id)

        markup = await back_to_keyboard(lang)
        if lang == "uz":
            await message.answer("Iltimos ismingizni kiriting 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Пожалуйста, введите ваше имя 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Please enter your name 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_name")
    else:
        markup =await language_keyboard()
        await message.answer(f"Kerakli tilni tanlang 👇\nChoose the language you need 👇\nВыберите нужный язык 👇", 
                            reply_markup=markup, protect_content=True)
        await state.set_state("get_lang")


@dp.message_handler(state="get_name", content_types=types.ContentTypes.TEXT)
async def get_name(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)            
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup =await language_keyboard()
        await message.answer(f"Kerakli tilni tanlang 👇\nChoose the language you need 👇\nВыберите нужный язык 👇", 
                            reply_markup=markup, protect_content=True)
        await state.set_state("get_lang")
    else:         
        user = await get_user(message.from_user.id)
        user.name = message.text
        user.save()
        markup = await phone_keyboard(lang)
        if lang == "uz":
            await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_phone_number")


@dp.message_handler(content_types=types.ContentTypes.CONTACT, state="get_phone_number")
async def get_phone(message: types.Message, state: FSMContext):
    if message.contact:
        phone = message.contact.phone_number[0:]
        user = await get_user(message.from_user.id)
        user.new_phone = phone
        otp = await generateOTP()
        await send_sms(otp=otp, phone=phone)
        user.otp = otp
        user.save()
        print(user.otp)
        lang = await get_lang(message.from_user.id)
        keyboard = await back_to_keyboard(lang)
        if lang == "uz":
            await message.answer(text=f"<b>{user.new_phone}</b> raqamiga yuborilgan tasdiqlash kodini kiriting", parse_mode='HTML', reply_markup=keyboard)
        if lang == "ru":
            await message.answer(text=f"Введите код подтверждения, отправленный на номер <b>{user.new_phone}</b>.", parse_mode='HTML', reply_markup=keyboard)
        if lang == "en":
            await message.answer(text=f"Enter the verification code sent to <b>{user.new_phone}</b>", parse_mode='HTML', reply_markup=keyboard)
        await state.set_state("get_otp")
    

@dp.message_handler(content_types=types.ContentTypes.TEXT, state="get_phone_number")
async def get_phone(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if "⬅️️" in message.text:
        user = await get_user(message.from_id)
        if user is not None:
            lang = await get_lang(message.from_user.id)
            if user.phone is not None:
                markup = await user_menu(lang)
                if lang == "uz":
                    await message.answer("Botimizga xush kelibsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
                elif lang == "en":
                    await message.answer("Welcome to our bot. Choose the section you want 👇", reply_markup=markup, protect_content=True)
                elif lang == "ru":
                    await message.answer("Добро пожаловать в наш бот. Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
                await state.set_state("get_category")
            else:
                markup = await phone_keyboard(lang)
                if lang == "uz":
                    await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing👇", reply_markup=markup, protect_content=True)
                elif lang == "en":
                    await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
                elif lang == "ru":
                    await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером👇", reply_markup=markup, protect_content=True)
                await state.set_state("get_phone_number")            
        else:
            markup =await language_keyboard()
            await message.answer(f"Assalomu alaykum, {message.from_user.first_name}👋. \nKerakli tilni tanlang 👇\n\nHello, {message.from_user.first_name}👋. \nChoose the language you need 👇\n\nЗдравствуйте, {message.from_user.first_name}👋. \nВыберите нужный язык 👇", 
                                reply_markup=markup, protect_content=True)
            await state.set_state("get_lang")
    else:
        if isValid(message.text):
            phone = message.text
            user = await get_user(message.from_user.id)
            user.new_phone = phone
            otp = await generateOTP()
            await send_sms(otp=otp, phone=phone)
            user.otp = otp
            user.save()
            print(user.otp)
            keyboard = await back_to_keyboard(lang)
            if lang == "uz":
                await message.answer(text=f"<b>{user.new_phone}</b> raqamiga yuborilgan tasdiqlash kodini kiriting", parse_mode='HTML', reply_markup=keyboard)
            if lang == "en":
                await message.answer(text=f"Введите код подтверждения, отправленный на номер <b>{user.new_phone}</b>.", parse_mode='HTML', reply_markup=keyboard)
            if lang == "ru":
                await message.answer(text=f"Enter the verification code sent to <b>{user.new_phone}</b>", parse_mode='HTML', reply_markup=keyboard)
            await state.set_state("get_otp")
        else:
            markup = await phone_keyboard(lang)
            if lang == "uz":
                await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing👇", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_phone_number")            
        

@dp.message_handler(lambda message: message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"], state="get_otp")
async def get_phone(message: types.Message, state: FSMContext):
    user = await get_user(message.from_user.id)
    lang = user.lang
    markup = await phone_keyboard(lang)
    if lang == "uz":
        await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing👇", reply_markup=markup, protect_content=True)
    elif lang == "en":
        await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
    elif lang == "ru":
        await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером👇", reply_markup=markup, protect_content=True)
    await state.set_state("get_phone_number")            


@dp.message_handler(content_types=types.ContentTypes.TEXT, state="get_otp")
async def get_phone(message: types.Message, state: FSMContext):
    user = await get_user(message.from_user.id)
    lang = user.lang
    if message.text == user.otp:
        user.phone = user.new_phone
        user.save()
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Botimizga xush kelighjk,bsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Welcome to our bot. Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Добро пожаловать в наш бот. Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    else:
        lang = await get_lang(message.from_user.id)
        markup = await back_to_keyboard(lang)
        if lang == "uz":
            await message.answer("⚠️ Yuborilgan tasdiqlash kodi xato. Qayta urinib ko'ring", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("⚠️ The verification code sent is incorrect. Try again", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("⚠️ Присланный проверочный код неверный. Попробуйте еще раз", reply_markup=markup, protect_content=True)
        await state.set_state("get_otp")


@dp.message_handler(state="get_category", content_types=[types.ContentTypes.TEXT, types.ContentTypes.DOCUMENT], commands=["import", "export", "settings", "contract", "customs", "cargo", "warehouse", "postal", "certification", "code", "contactus", "feedback", "address", "exchange", "library"])
@dp.message_handler(state="get_category", content_types=[types.ContentTypes.TEXT, types.ContentTypes.DOCUMENT])
async def get_service_category(message: types.Message, state: FSMContext):
    await message.answer(message.document.file_id)
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    user = await get_user(message.from_user.id)
    cmd = str(message.get_command())
    await state.update_data(state=message.text)
    message_id = int(message.message_id) + 1
    category = await get_category_by_name(message.text)
    if category is not None and category != []:
        user.interests.add(category)
        user.save()
    if message.text in ["Sozlamalar ⚙️", "Настройки ⚙️", "Settings ⚙️"] or cmd == "/settings":
        markup = await settings_keyboard(lang)
        if lang == "uz":
            await message.answer(text="Kerakli buyruqni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer(text="Choose the command you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer(text="Выберите нужную команду 👇", reply_markup=markup, protect_content=True)
        await state.set_state("settings")
    elif message.text in ["Библиотека 📚", "Kutubxona 📚", "Library 📚"] or cmd == "/library":
        markup = await library_keyboard(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Please select the desired section 👇", reply_markup=markup, protect_content=True)
        await state.set_state("library")
        # doc = open("./qaror.pdf", 'rb')
        # markup = await user_menu(lang)
        # await message.answer_document(document=doc, reply_markup=markup, protect_content=True)
    elif message.text in ["Valyutalar kursi 💳", "Exchange rates 💳", "Курсы обмена валюты 💳"] or cmd == "/exchange":
        markup = await user_menu(lang)
        kurslar = valyuta_kurslari()
        if lang == "uz":
            await message.answer(text=kurslar)
            await message.answer(text="Kerakli buyruqni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer(text=kurslar)
            await message.answer(text="Choose the command you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer(text=kurslar)
            await message.answer(text="Выберите нужную команду 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    elif message.text in ["Import 🚚", "Импорт 🚚"]  or cmd == "/import":
        if user.full:
            if lang == "uz":
                await message.answer("Maxsulot nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter the product name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название продукта 👇", reply_markup=back_key)
            await state.set_state("import_product_name")
        else:
            if lang == "uz":
                await message.answer("Korxonangiz nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter your company name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название вашей компании 👇", reply_markup=back_key)
            await state.set_state("get_company_name")
    if message.text in ["Export 🚛", "Экспорт 🚛"] or cmd == "/export":
        if user.full:
            if lang == "uz":
                await message.answer("Maxsulot nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter the product name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название продукта 👇", reply_markup=back_key)
            await state.set_state("export_product_name")   
        else:
            if lang == "uz":
                await message.answer("Korxonangiz nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter your company name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название вашей компании 👇", reply_markup=back_key)
            await state.set_state("get_company_name")
    if message.text in ["Contract 🗂", "Kontrakt 🗂", "Контракт 🗂"]  or cmd == "/contract":
        if user.full:
            markup = await kontrakt_keyboard(lang)
            if lang == "uz":
                await message.answer("Kerakli xizmat turini tanlang 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer("Choose the type of service you need 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer("Выберите нужный вам вид услуги 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_contract_service")
        else:
            if lang == "uz":
                await message.answer("Korxonangiz nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter your company name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название вашей компании 👇", reply_markup=back_key)
            await state.set_state("get_company_name")
    if message.text in ["TIF bojxona ro'yxati 🏫", "TIF customs list 🏫", "Тифозный таможенный список 🏫"] or cmd == "/customs":
        back_key = await back_to_keyboard(lang)
        markup = await region_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_customs_region")
    if message.text in ["Yuk xizmatlari 📦", "Freight services 📦", "Грузовые услуги 📦"] or cmd == "/cargo":
        markup = await freight_keyboard(lang)
        back_key = await back_to_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli xizmat turini tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Choose the type of service you need 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный вам вид услуги 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_freight_service")
    if message.text in ["Omborlar ro'yxati 🏢", "Warehouse list 🏢", "Список складов 🏢"] or cmd == "/warehouse":
        back_key = await back_to_keyboard(lang)
        markup = await region_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_region")                                      
    if message.text in ["Vagon kuzatish 🚃", "Track carriage 🚃", "Отслеживать вагон 🚃"] or cmd == "/track":
        pass
    if message.text in ["Pochta xizmati 📨", "Postal service 📨", "Почтовая служба 📨"] or cmd == "/postal":
        back_key = await back_to_keyboard(lang)
        markup = await region_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_post_region")                                      
    if message.text in ["Sertifikatlash 📑", "Certification 📑", "Сертификация 📑"] or cmd == "/certification":
        back_key = await back_to_keyboard(lang)
        await state.update_data(page=1)
        max_data = await get_sertification_count()
        # markup = await sertification_keyboard(lang=lang, page=1)
        markup = InlineKeyboardMarkup().add(InlineKeyboardButton(text='Sertifikatlash 📑', callback_data='sert'),
                                            InlineKeyboardButton(text='Sug\'urtalash 📃', callback_data='sug'))
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer(f"Kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired Certificate Authority 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_sertification")
    if message.text in ["TN VED Kodi 🆔", "HS CODE 🆔", "Код ТНВЭД 🆔"] or cmd == "/code":
        back_key = await back_to_keyboard(lang)
        markup = await tnved_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli kodni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired code 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный код 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_tnved")


@dp.callback_query_handler(state="get_tnved")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        await call.message.delete()
        markup = await user_menu(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text="Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text="Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text="Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    else:
        tenved = await get_tenved_id(call.data)
        markup = await back(lang)
        if lang == "uz":
            await call.message.edit_text(f"{tenved.kod} - {tenved.description_uz}", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(f"{tenved.kod} - {tenved.description_en}", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(f"{tenved.kod} - {tenved.description_uz}", reply_markup=markup)
        await state.set_state("tenved")


@dp.callback_query_handler(state="get_sertification")
async def get_pod(c: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(c.from_user.id)
    if c.data == "sert":
        await state.update_data(page=1)
        max_data = await get_sertification_count()
        markup = await sertification_keyboard(lang=lang, page=1)
        if lang == "uz":
            await c.message.edit_text(f"Jami ma'lumotlar {max_data} ta. Kerakli Sertifikatlar idorasini tanlang 👇", reply_markup=markup)
        if lang == "en":
            await c.message.answer(f"Total data in {max_data}. Select the desired Certificate Authority 👇", reply_markup=markup)
        if lang == "ru":
            await c.message.answer(f"Всего данных на {max_data}. Выберите нужный центр сертификации 👇", reply_markup=markup)
        await state.set_state("sertification")

@dp.callback_query_handler(state="tenved")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await tnved_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("kodni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Select the desired code 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный код 👇", reply_markup=markup)
        await state.set_state("get_tnved")


@dp.message_handler(state="library", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Please select the desired section 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    else:
        doc = open("./qaror.pdf", 'rb')
        markup = await library_keyboard(lang)
        await message.answer_document(document=doc, reply_markup=markup, protect_content=True)


@dp.callback_query_handler(state="sertification")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    state_data = await state.get_data()
    command = call.data
    this_page = state_data['page']
    page = 0
    if command == "back":
        await call.message.delete()
        markup = await user_menu(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text="Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text="Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text="Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    elif command == "next_page":
        max_page = await get_sertification_max_page()
        if int(this_page) == max_page:
            page = 1
        else:
            page = int(this_page) + 1
        markup = await sertification_keyboard(page=page, lang=lang)
        await call.message.edit_reply_markup(reply_markup=markup, protect_content=True)
        await state.update_data(page=page)
    elif command == "last_page":
        max_page = await get_sertification_max_page()
        if int(this_page) == 1:
            page = max_page
        else:
            page = int(this_page) - 1
        markup = await sertification_keyboard(page=page, lang=lang)
        await call.message.edit_reply_markup(reply_markup=markup, protect_content=True)
        await state.update_data(page=page)
    else:
        sert = await get_sertification(command)
        markup = await back(lang)
        if lang == "uz":
            await call.message.edit_text(text=f"{sert.legalname_uz}\n\n{sert.address_uz}\n\n{sert.contacts}\n\n", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text=f"{sert.legalname_en}\n\n{sert.address_en}\n\n{sert.contacts}\n\n", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text=f"{sert.legalname_ru}\n\n{sert.address_ru}\n\n{sert.contacts}\n\n", reply_markup=markup)
        await state.set_state("get_sertificate")


@dp.callback_query_handler(state="get_sertificate")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    state_data = await state.get_data()
    command = call.data
    this_page = state_data['page']
    max_data = await get_sertification_count()
    page = 0
    if command == "back":
        markup = await sertification_keyboard(lang=lang, page=int(this_page))
        if lang == "uz":
            await call.message.edit_text(text=f"Jami ma'lumotlar {max_data} ta. Kerakli Sertifikatlash idorasini tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text="Total data in {max_data}. Select the desired Certificate Authority 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text=f"Всего данных на {max_data}. Выберите нужный центр сертификации 👇", reply_markup=markup)
        await state.set_state('sertification')


@dp.callback_query_handler(state="get_customs_region")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        await call.message.delete()
        markup = await user_menu(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text="Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text="Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text="Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    else:
        region = await get_region(call.data)
        await state.update_data(region=call.data)
        markup = await customs_keyboard(lang=lang, region=call.data)
        if lang == "uz":
            await call.message.edit_text(text=f"{region.name_uz}dagi TIF bojxonalari 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text="TIF customs offices in {region.name_en} 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text=f"Таможенные посты TIF в {region.name_ru} 👇", reply_markup=markup)
        await state.set_state("get_tif")


@dp.callback_query_handler(state="get_tif")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text(text="Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text="Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text="Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_customs_region") 
    else:
        await call.message.delete()
        customs = await get_one_customs(call.data)
        await bot.send_location(chat_id=call.from_user.id, longitude=customs.longitude, latitude=customs.latitude)
        markup = await back(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text=f"{customs.name_uz}\n{customs.contact}", reply_markup=markup, protect_content=True)
        if lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text=f"{customs.name_en}\n{customs.contact}", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text=f"{customs.name_ru}\n{customs.contact}", reply_markup=markup, protect_content=True)
        await state.set_state("custom")


@dp.callback_query_handler(state="custom")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        data = await state.get_data()
        region_id = data["region"]
        region = await get_region(region_id)
        markup = await customs_keyboard(lang=lang, region=region_id)
        if lang == "uz":
            await call.message.edit_text(text=f"{region.name_uz}dagi TIF bojxonalari 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text="TIF customs offices in {region.name_en} 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text=f"Таможенные посты TIF в {region.name_ru} 👇", reply_markup=markup)
        await state.set_state("get_tif")
        
        
@dp.callback_query_handler(state="get_post_region")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        await call.message.delete()
        markup = await user_menu(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text="Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text="Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text="Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    else:
        counts = await get_region_posts(call.data)
        await state.update_data(region_id=call.data)
        text = ""
        markup = await posts_keyboard(region=call.data, lang=lang)
        region = await get_region(call.data)
        if lang == "uz":
            text += f"{region.name_uz} viloyatida {counts} ta omborxona mavjud. Ular:"
        if lang == "en":
            text += f"There are {counts} warehouses in {region.name_en}. They are:"
        if lang == "ru":
            text += f"В регионе {region.name_ru} есть {counts} складов. Они есть:"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.set_state("get_post_service")


@dp.callback_query_handler(state="get_post_service")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text(text="Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text="Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text="Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_post_region") 
    else:
        await call.message.delete()
        post = await get_post(call.data)
        text = ""
        markup = await back_to_keyboard(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text=f"{post.legalname_uz}\n\n{post.description_uz}\n\nManzil: {post.address_uz}\n\nKontakt: {post.contacts}", reply_markup=markup, protect_content=True)
        if lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text=f"{post.legalname_en}\n\n{post.description_en}\n\nAddress: {post.address_uz}\n\nContacts: {post.contacts}", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text=f"{post.legalname_ru}\n\n{post.description_ru}\n\nАдрес: {post.address_uz}\n\nКонтактное лицо: {post.contacts}", reply_markup=markup, protect_content=True)
        await state.set_state("post")


@dp.message_handler(state="post", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        data = await state.get_data()
        region_id = data["region_id"]
        counts = await get_region_posts(region_id)
        await state.update_data(region_id=region_id)
        text = ""
        markup = await posts_keyboard(region=region_id, lang=lang)
        region = await get_region(region_id)
        if lang == "uz":
            text += f"{region.name_uz} viloyatida {counts} ta omborxona mavjud. Ular:"
        if lang == "en":
            text += f"There are {counts} warehouses in {region.name_en}. They are:"
        if lang == "ru":
            text += f"В регионе {region.name_ru} есть {counts} складов. Они есть:"
        await message.answer(text=text, reply_markup=markup, protect_content=True)
        await state.set_state("get_post_service")
                                                     


@dp.callback_query_handler(state="get_region")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        await call.message.delete()
        markup = await user_menu(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text="Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text="Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text="Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    else:
        counts = await get_region_wearhouses(call.data)
        await state.update_data(region_id=call.data)
        text = ""
        markup = await wearhouses_keyboard(region_id=call.data, lang=lang)
        region = await get_region(call.data)
        if lang == "uz":
            text += f"{region.name_uz} viloyatida {counts} ta omborxona mavjud. Ular:"
        if lang == "en":
            text += f"There are {counts} warehouses in {region.name_en}. They are:"
        if lang == "ru":
            text += f"В регионе {region.name_ru} есть {counts} складов. Они есть:"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.set_state("get_wearhouse")


@dp.callback_query_handler(state="get_wearhouse")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text(text="Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text(text="Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text(text="Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_region") 
    else:
        await call.message.delete()
        wearhouse = await get_wearhouse(call.data)
        await bot.send_location(chat_id=call.from_user.id, longitude=wearhouse.longitude, latitude=wearhouse.latitude)
        text = ""
        markup = await back_keyboard(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text=f"{wearhouse.name_uz}\n\n{wearhouse.description_uz}", reply_markup=markup, protect_content=True)
        if lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text=f"{wearhouse.name_en}\n\n{wearhouse.description_en}", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text=f"{wearhouse.name_ru}\n\n{wearhouse.description_ru}", reply_markup=markup, protect_content=True)
        await state.set_state("wearhouse")


@dp.message_handler(state="wearhouse", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        data = await state.get_data()
        region_id = data["region_id"]
        counts = await get_region_wearhouses(region_id)
        await state.update_data(region_id=region_id)
        text = ""
        markup = await wearhouses_keyboard(region_id=region_id, lang=lang)
        region = await get_region(region_id)
        if lang == "uz":
            text += f"{region.name_uz} viloyatida {counts} ta omborxona mavjud. Ular:"
        if lang == "en":
            text += f"There are {counts} warehouses in {region.name_en}. They are:"
        if lang == "ru":
            text += f"В регионе {region.name_ru} есть {counts} складов. Они есть:"
        await message.answer(text=text, reply_markup=markup, protect_content=True)
        await state.set_state("get_wearhouse")
                                                     

@dp.message_handler(state="get_freight_service", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')


@dp.message_handler(state="get_location", content_types=types.ContentTypes.LOCATION)
async def get_locations(message: types.Message, state: FSMContext):
    location = message.location
    lang = await get_lang(message.from_user.id)
    user_location = {"lon": location.longitude, "lat": location.latitude}
    data = await get_adresses()
    closect_wearhouse = closest(data=data, location=user_location)
    wearhouse = await get_wearhouse(closect_wearhouse['id'])
    if lang == "uz":
        await message.answer(text=f"Sizga eng yaqin bo'lgan omborxona\n\n {wearhouse.name_uz}\n\n{wearhouse.description_uz}")
    if lang == "en":
        await message.answer(text=f"Your nearest warehouse\n\n{wearhouse.name_en}\n\n{wearhouse.description_en}")
    if lang == "ru":
        await message.answer(text=f"Ближайший к вам склад\n\n{wearhouse.name_ru}\n\n{wearhouse.description_ru}")
    await message.answer_location(longitude=wearhouse.longitude, latitude=wearhouse.latitude)


@dp.message_handler(state="get_location", content_types=types.ContentTypes.TEXT)
async def get_location(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')


@dp.message_handler(state="get_region", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')


@dp.message_handler(state="get_wearhouse", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await region_keyboard(lang)
        if lang == "uz":
            await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_region")                                      


@dp.callback_query_handler(state="get_freight_service")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        await call.message.delete()
        markup = await user_menu(lang)
        if lang == "uz":
            await bot.send_message(chat_id=call.from_user.id, text="Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await bot.send_message(chat_id=call.from_user.id, text="Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await bot.send_message(chat_id=call.from_user.id, text="Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    if command == "loader_equipment":
        markup = await loader_equipment_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text(text="Iltimos xizmat turini tanlang 👇", reply_markup=markup)
        elif lang == "en":
            await call.message.edit_text(text="Please select the type of service 👇", reply_markup=markup)
        elif lang == "ru":
            await call.message.edit_text(text="Пожалуйста, выберите тип услуги 👇", reply_markup=markup)
        await state.set_state('get_equipment_type')
    if command == "loader_service":
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_loader_region")
    if command == "shipping":
        text = ""
        markup = await logistics_keyboard(lang)
        if lang == 'uz':
            text += f"Kerakli xizmat turni tanlang 👇"
        if lang == 'en':
            text += f"Choose the type of service you need 👇"
        if lang == 'ru':
            text += f"Выберите тип услуги, который вам нужен 👇"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.set_state("get_logistics_service")


@dp.callback_query_handler(state="get_logistics_service")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await freight_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli xizmat turini tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Choose the type of service you need 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный вам вид услуги 👇", reply_markup=markup)
        await state.set_state("get_freight_service")
    elif command == "internal":
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_logistics_region")


@dp.callback_query_handler(state="get_logistics_region")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        text = ""
        markup = await logistics_keyboard(lang)
        if lang == 'uz':
            text += f"Kerakli xizmat turni tanlang 👇"
        if lang == 'en':
            text += f"Choose the type of service you need 👇"
        if lang == 'ru':
            text += f"Выберите тип услуги, который вам нужен 👇"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.set_state("get_logistics_service")
    else:
        await state.update_data(region=call.data)
        text = ""
        markup = await tonna_keyboard(lang)
        if lang == 'uz':
            text += f"Kerakli og'irlikni tanlang(tonnalarda) tanlang 👇"
        if lang == 'en':
            text += f"Select the desired weight (in tons) 👇"
        if lang == 'ru':
            text += f"Выберите нужный вес (в тоннах) 👇"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.set_state("get_logistics_tonna")                       


@dp.callback_query_handler(state="get_logistics_tonna")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_logistics_region")
    else:
        state_data = await state.get_data()
        region = state_data['region']
        services = await get_by_tonna(tonna=call.data, region=region)
        await state.update_data(page=1, tonna=call.data)
        objects = await logistic_pagination(page=1, data=services)
        text = ""
        markup = await pagination_keyboard(lang)
        for object in objects:
            if lang == "uz":
                text += f"Turi: {object.type}       Hudud: {object.region.name_uz}\nTelefon: {object.phone}\nTonna: {object.tonna} t.\n\n"
            if lang == "en":
                text += f"Type: {object.type}       Region: {object.region.name_uz}\nPhone: {object.phone}\nTonna: {object.tonna} t.\n\n"
            if lang == "ru":
                text += f"Тип: {object.type}        Регион: {object.region.name_uz}\nТелефон: {object.phone}\nТонна: {object.tonna} т.\n\n"
        await call.message.edit_text(text=text, reply_markup=markup.set_state("logistic_service"))


@dp.callback_query_handler(state="logistic_service")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    state_data = await state.get_data()
    command = call.data
    this_page = state_data['page']
    page = 0
    region = state_data['region']
    tonna = state_data['tonna']
    if command == "back":
        text = ""
        markup = await tonna_keyboard(lang)
        if lang == 'uz':
            text += f"Kerakli og'irlikni tanlang(tonnalarda) tanlang 👇"
        if lang == 'en':
            text += f"Select the desired weight (in tons) 👇"
        if lang == 'ru':
            text += f"Выберите нужный вес (в тоннах) 👇"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.set_state("get_logistics_tonna")                       
    elif command == "next_page":
        objects = await get_by_tonna(tonna=tonna, region=region)
        max_page = len(objects) // 15 + 1
        if int(this_page) == max_page:
            page = 1
        else:
            page = int(this_page) + 1
        objects = await logistic_pagination(page=page, data=objects)
        text = ""
        markup = await pagination_keyboard(lang)
        for object in objects:
            if lang == "uz":
                text += f"Turi: {object.type}       Hudud: {object.region.name_uz}\nTelefon: {object.phone}\nTonna: {object.tonna} t.\n\n"
            if lang == "en":
                text += f"Type: {object.type}       Region: {object.region.name_uz}\nPhone: {object.phone}\nTonna: {object.tonna} t.\n\n"
            if lang == "ru":
                text += f"Тип: {object.type}        Регион: {object.region.name_uz}\nТелефон: {object.phone}\nТонна: {object.tonna} т.\n\n"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.update_data(page=page)
    elif command == "last_page":
        objects = await get_by_tonna(tonna=tonna, region=region)
        max_page = len(objects) // 15 + 1
        if int(this_page) == 1:
            page = max_page
        else:
            page = int(this_page) - 1
        objects = await logistic_pagination(page=page, data=objects)
        text = ""
        markup = await pagination_keyboard(lang)
        for object in objects:
            if lang == "uz":
                text += f"Turi: {object.type}       Hudud: {object.region.name_uz}\nTelefon: {object.phone}\nTonna: {object.tonna} t.\n\n"
            if lang == "en":
                text += f"Type: {object.type}       Region: {object.region.name_uz}\nPhone: {object.phone}\nTonna: {object.tonna} t.\n\n"
            if lang == "ru":
                text += f"Тип: {object.type}        Регион: {object.region.name_uz}\nТелефон: {object.phone}\nТонна: {object.tonna} т.\n\n"
        await call.message.edit_text(text=text, reply_markup=markup)
        await state.update_data(page=page)

        
@dp.message_handler(state="loader_service", content_types=types.ContentTypes.TEXT)
async def loader_service(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    message_id = int(message.message_id) + 1
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        back_key = await back_to_keyboard(lang)
        markup = await freight_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli xizmat turini tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Choose the type of service you need 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=back_key)
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный вам вид услуги 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_freight_service")
    

@dp.callback_query_handler(state="get_equipment_type")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await freight_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli xizmat turini tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Choose the type of service you need 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный вам вид услуги 👇", reply_markup=markup)
        await state.set_state("get_freight_service")
    else:
        await state.update_data(equipment_type=command)
        markup = await region_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli viloyatni tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Select the desired region 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный регион 👇", reply_markup=markup)
        await state.set_state("get_equipment_region")
        
        
@dp.callback_query_handler(state="get_loader_region")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await freight_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text("Kerakli xizmat turini tanlang 👇", reply_markup=markup)
        if lang == "en":
            await call.message.edit_text("Choose the type of service you need 👇", reply_markup=markup)
        if lang == "ru":
            await call.message.edit_text("Выберите нужный вам вид услуги 👇", reply_markup=markup)
        await state.set_state("get_freight_service")
    else:
        await call.message.delete()
        text = ""
        loaders = await get_loaders(call.data)
        if lang == "uz":
            text = "Ush hududda xizmat ko'rsatuvchi yuklovchilar 👇"
        if lang == "en":
            text = "Loaders serving in that area 👇"
        if lang == "ru":
            text = "Грузчики, работающие в этом районе 👇"
        i = 1
        for loader in loaders:
            if lang == "uz":
                text += f"\n\n{i})Turi: {loader.name_uz}\n   Telefon: {loader.phone}\n"
            if lang == "en":
                text += f"\n\n{i})Type: {loader.name_en}\n  Phone: {loader.phone}\n"
            if lang == "ru":
                text += f"\n\n{i})Тип: {loader.name_ru}\n  Телефон: {loader.phone}\n"
            i += 1
        markup = await back_keyboard(lang)
        try:
            await bot.send_message(chat_id=call.from_user.id, text=text, reply_markup=markup, protect_content=True)
        except:
            await bot.send_message(chat_id=call.from_user.id, text=text[:4000], reply_markup=markup, protect_content=True)
        await state.set_state("loaders")

        
@dp.message_handler(state="loaders", content_types=types.ContentTypes.TEXT)
async def loader_service(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    message_id = int(message.message_id) + 1
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        back_key = await back_to_keyboard(lang)
        markup = await region_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=ReplyKeyboardRemove())
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=ReplyKeyboardRemove())
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=ReplyKeyboardRemove())
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_loader_region")
    

@dp.callback_query_handler(state="get_equipment_region")
async def get_tif(call: types.CallbackQuery, state: FSMContext):
    lang = await get_lang(call.from_user.id)
    command = call.data
    if command == "back":
        markup = await loader_equipment_keyboard(lang)
        if lang == "uz":
            await call.message.edit_text(text="Iltimos xizmat turini tanlang 👇", reply_markup=markup)
        elif lang == "en":
            await call.message.edit_text(text="Please select the type of service 👇", reply_markup=markup)
        elif lang == "ru":
            await call.message.edit_text(text="Пожалуйста, выберите тип услуги 👇", reply_markup=markup)
        await state.set_state('get_equipment_type')
    else:
        data = await state.get_data()
        region_id = command
        equipment_type = data["equipment_type"]
        equipments = await get_loader_equipments(type=equipment_type, region=region_id)
        text = ""
        i = 1
        if equipments and equipments is not None:
            await call.message.delete()
            for equipment in equipments:
                if lang == "uz":
                    text += f"{i}) {equipment.name_uz}.\n\n"     
                if lang == "en":
                    text += f"{i}) {equipment.name_en}.\n\n"     
                if lang == "ru":
                    text += f"{i}) {equipment.name_ru}.\n\n"   
                i += 1
            back_key = await back_keyboard(lang)
            await bot.send_message(chat_id=call.from_user.id, text=text, reply_markup=back_key)
            await state.set_state('equipments')
        else:
            markup = await loader_equipment_keyboard(lang)
            if lang == "uz":
                await call.message.edit_text(text="🚫 Ma'lumotlar topilmadi. \n\nIltimos xizmat turini tanlang 👇", reply_markup=markup)
            elif lang == "en":
                await call.message.edit_text(text="🚫 No data found. \n\nPlease select the type of service 👇", reply_markup=markup)
            elif lang == "ru":
                await call.message.edit_text(text="🚫 Данные не найдены. \n\nПожалуйста, выберите тип услуги 👇", reply_markup=markup)
            await state.set_state('get_equipment_type')

        
@dp.message_handler(state="equipments", content_types=types.ContentTypes.TEXT)
async def loader_service(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    message_id = int(message.message_id) + 1
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        back_key = await back_to_keyboard(lang)
        markup = await loader_equipment_keyboard(lang)
        if lang == "uz":
            await message.answer(".", reply_markup=ReplyKeyboardRemove())
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(".", reply_markup=ReplyKeyboardRemove())
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(".", reply_markup=ReplyKeyboardRemove())
            await bot.delete_message(chat_id=message.from_id, message_id=message_id)
            await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_loader_region")
    

              
@dp.message_handler(state="get_tif", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
 
 
@dp.message_handler(state="get_contract_service", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    get_phone = await get_phone_keyboard(lang)
    if message.text in  ["Conclusion of an import/export contract", "Registration under YeEISVO", "Solving contract problems", "Заключение импортно-экспортного контракта", "Регистрация в ЕЭИСВО", "Решение проблем с контрактами", "Import/eksport shartnoma tuzish", "YeEISVO bo'yicha ro'yxatdan o'tish", "Kontraktdagi muammolarni hal qilish"]:
        await state.update_data(contract_type=message.text)
        if lang == "uz":
            await message.answer(f"{message.text} bo'yicha sizga qanday yordam bera olamiz? 1 knopka yordamida qo'ng'iroq buyurtma qiling.👇", reply_markup=get_phone)
        if lang == "en":
            await message.answer(f"How can we help you with {message.text}? Order a call using 1 button.👇", reply_markup=get_phone)
        if lang == "ru":
            await message.answer(f"Как мы можем помочь вам с {message.text}? Заказать звонок с помощью 1 кнопки.👇", reply_markup=get_phone)
        await state.set_state("get_phone_order")


@dp.message_handler(state="get_phone_order", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await kontrakt_keyboard(lang)
        if lang == "uz":
            await message.answer("Kerakli xizmat turini tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer("Choose the type of service you need 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer("Выберите нужный вам вид услуги 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_contract_service")
    if message.text in ["Qo'ng'iroq buyurtma qilish", "Order a call", "Заказ звонка"]:
        data = await state.get_data()
        service = data["contract_type"]
        markup = await user_menu(lang)
        user = await get_user(message.from_user.id)
        await bot.send_message(chat_id=-838866316, text=f"{user.name}\n\nTelefon: {user.phone}\n\nService{service}")
        if lang == "uz":
            await message.answer(f"Hodimimiz siz bilan 24 soat ichida bog'lanadi", reply_markup=markup, protect_content=True)
            await message.answer(f"Bosh menyu. Kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(f"Our staff will contact you within 24 hours", reply_markup=markup, protect_content=True)
            await message.answer(f"Main menu. Select the desired section 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(f"Наши сотрудники свяжутся с вами в течение 24 часов", reply_markup=markup, protect_content=True)
            await message.answer(f"Главное меню. Выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    if message.text in ["Savol qoldirish", "Leave a question", "Оставить вопрос"]:
        if lang == "uz":
            await message.answer("Savolingizni qoldiring 👇", reply_markup=ReplyKeyboardRemove())
        if lang == "en":
            await message.answer("Leave your question 👇", reply_markup=ReplyKeyboardRemove())
        if lang == "ru":
            await message.answer("Оставьте свой вопрос 👇", reply_markup=ReplyKeyboardRemove())
        await state.set_state("get_contract_question")


@dp.message_handler(state="get_contract_question", content_types=types.ContentTypes.TEXT)
async def get_contract_question(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    await message.forward(chat_id=-838866316)
    markup = await user_menu(lang)
    if lang == "uz":
        await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
    elif lang == "en":
        await message.answer("Choose the section you want 👇", reply_markup=markup, protect_content=True)
    elif lang == "ru":
        await message.answer("Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
    await state.set_state('get_category')
             

@dp.message_handler(state="get_company_name", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Botimizga xush kelibsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Welcome to our bot. Choose the section you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Добро пожаловать в наш бот. Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_category')
    else:
        markup = await product_categories(lang)
        user = await get_user(message.from_user.id)
        user.company = message.text
        user.save()
        if lang == "uz":
            await message.answer("Firmangiz kategoriyasini tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Select the category of your company 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Выберите категорию вашей компании 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_product_category')


@dp.message_handler(state="get_product_category", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    user = await get_user(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        if lang == "uz":
            await message.answer("Firmangiz nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter your company name 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название вашей компании 👇", reply_markup=back_key)
        await state.set_state("get_company_name")
    else:
        category = await get_product_category_by_name(message.text)
        if category is not None:
            user.product_category = category
            user.save()
            markup = await get_company_monthly(lang)
            if lang == "uz":
                await message.answer("Firmangiz oylik aylanmasini (Tonnalarda) tanlang 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer("Select your company's monthly turnover (in Tons) 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer("Выберите месячный оборот вашей компании (в тоннах) 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_company_monthly")
        else:
            markup = await product_categories(lang)
            if lang == "uz":
                await message.answer("Firmangiz kategoriyasini tanlang 👇", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer("Select the category of your company 👇", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer("Выберите категорию вашей компании 👇", reply_markup=markup, protect_content=True)
            await state.set_state('get_product_category')
          

@dp.message_handler(state="get_company_monthly", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    user = await get_user(message.from_user.id)
    message_id = int(message.message_id) + 1
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await product_categories(lang)
        if lang == "uz":
            await message.answer("Firmangiz kategoriyasini tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Select the category of your company 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Выберите категорию вашей компании 👇", reply_markup=markup, protect_content=True)
        await state.set_state('get_product_category')
    else:
        user = await get_user(message.from_user.id)
        user.monthly = message.text
        user.full = True
        user.save()
        data = await state.get_data()
        command = data ['state']
        if lang == "uz":
            await message.answer("Ma'lumotlar qabul qilindi ✅")
        if lang == "en":
            await message.answer("The information has been accepted ✅")
        if lang == "ru":
            await message.answer("Данные получены успешно ✅")
        if command in ["Sozlamalar", "Настройки", "Settings"]:
            markup = await settings_keyboard(lang)
            if lang == "uz":
                await message.answer(text="Kerakli buyruqni tanlang 👇", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer(text="Choose the command you want 👇", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer(text="Выберите нужную команду 👇", reply_markup=markup, protect_content=True)
            await state.set_state("settings")
        elif command in ["Import", "Импорт"]:
            if lang == "uz":
                await message.answer("Maxsulot nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter the product name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название продукта 👇", reply_markup=back_key)
            await state.set_state("import_product_name")
        elif command in ["Export", "Экспорт"]:
            if lang == "uz":
                await message.answer("Maxsulot nomini kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter the product name 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите название продукта 👇", reply_markup=back_key)
            await state.set_state("export_product_name")   
        elif command in ["Contract", "Kontrakt", "Контракт"]:
            markup = await kontrakt_keyboard(lang)
            if lang == "uz":
                await message.answer("Kerakli xizmat turini tanlang 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer("Choose the type of service you need 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer("Выберите нужный вам вид услуги 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_contract_service")
        elif command in ["TIF bojxona ro'yxati", "TIF customs list", "Тифозный таможенный список"]:
            back_key = await back_to_keyboard(lang)
            markup = await customs_keyboard(lang)
            if lang == "uz":
                await message.answer("TIF bojxona ro'yxati:", reply_markup=back_key)
                await message.answer("Kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer("TIF customs list:", reply_markup=back_key)
                await message.answer("Select the desired section 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer("Таможенный список брюшного тифа:", reply_markup=markup, protect_content=True)
                await message.answer("Выберите нужный раздел 👇", reply_markup=back_key)
            await state.set_state("get_tif") 
        elif command in ["Yuk xizmatlari", "Freight services", "Грузовые услуги"]:
            markup = await freight_keyboard(lang)
            back_key = await back_to_keyboard(lang)
            if lang == "uz":
                await message.answer(".", reply_markup=back_key)
                await bot.delete_message(chat_id=message.from_id, message_id=message_id)
                await message.answer("Kerakli xizmat turini tanlang 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer(".", reply_markup=back_key)
                await bot.delete_message(chat_id=message.from_id, message_id=message_id)
                await message.answer("Choose the type of service you need 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer(".", reply_markup=back_key)
                await bot.delete_message(chat_id=message.from_id, message_id=message_id)
                await message.answer("Выберите нужный вам вид услуги 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_freight_service")
        elif command in ["Omborlar ro'yxati", "Warehouse list", "Список складов"]:
            back_key = await back_to_keyboard(lang)
            markup = await region_keyboard(lang)
            if lang == "uz":
                await message.answer(".", reply_markup=back_key)
                await bot.delete_message(chat_id=message.from_id, message_id=message_id)
                await message.answer("Kerakli viloyatni tanlang 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer(".", reply_markup=back_key)
                await bot.delete_message(chat_id=message.from_id, message_id=message_id)
                await message.answer("Select the desired region 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer(".", reply_markup=back_key)
                await bot.delete_message(chat_id=message.from_id, message_id=message_id)
                await message.answer("Выберите нужный регион 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_region")                                      
        elif command in ["Eng yaqin manzillar", "Nearest addresses", "Самые близкие адреса"]:
            markup = await location_send(lang)
            if lang == "uz":
                await message.answer("Joylashuv manzilingizni jo'nating 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer("Please send your location address 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer("Отправьте свое местоположение 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_location")
        else:
            markup = await user_menu(lang)
            if lang == "uz":
                await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer("Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer("Please select the desired section 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_category")


@dp.message_handler(content_types=types.ContentTypes.TEXT, state="settings")
async def get_settings_message(message: types.Message, state:FSMContext):
    lang = await get_lang(message.from_user.id)
    if "⬅️" in  message.text:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Select the required button👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Выберите нужную кнопку👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    elif message.text in ["🔄 Tilni o'zgartirish", "🔄 Изменить язык", "🔄 Change language"]:
        if lang == "uz":
            markup = await language_keyboard()
            await message.answer(text="Tilni o'zgartirish ♻️\nKerakli tilni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            markup = await language_keyboard()
            await message.answer(text="Change language ♻️\nChoose the language you want 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            markup = await language_keyboard()
            await message.answer(text="Изменить язык ♻️\nВыберите нужный язык 👇", reply_markup=markup, protect_content=True)
        await state.set_state("set_lang")
    elif message.text in ["📞 Raqamni o'zgartirish", "📞 Изменить номер телефона", "📞 Change phone number"]:
        markup = await phone_keyboard(lang)
        if lang == "uz":
            await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_phone_number_settings")            


@dp.message_handler(content_types=types.ContentTypes.CONTACT, state="get_phone_number_settings")
async def get_phone(message: types.Message, state: FSMContext):
    if message.contact:
        phone = message.contact.phone_number[1:]
        user = await get_user(message.from_user.id)
        user.new_phone = phone
        otp = await generateOTP()
        await send_sms(otp=otp, phone=phone)
        user.otp = otp
        user.save()
        print(user.otp)
        lang = await get_lang(message.from_user.id)
        keyboard = await back_keyboard(lang)
        if lang == "uz":
            await message.answer(text=f"<b>{user.new_phone}</b> raqamiga yuborilgan tasdiqlash kodini kiriting", parse_mode='HTML', reply_markup=keyboard)
        if lang == "ru":
            await message.answer(text=f"Введите код подтверждения, отправленный на номер <b>{user.new_phone}</b>.", parse_mode='HTML', reply_markup=keyboard)
        if lang == "en":
            await message.answer(text=f"Enter the verification code sent to <b>{user.new_phone}</b>", parse_mode='HTML', reply_markup=keyboard)
        await state.set_state("get_otp_settings")
    

@dp.message_handler(content_types=types.ContentTypes.TEXT, state="get_phone_number_settings")
async def get_phone_settings(message: types.Message, state: FSMContext):
    if "⬅️" not in message.text:
        lang = await get_lang(message.from_user.id)
        if isValid(message.text):
            phone = message.text
            user = await get_user(message.from_user.id)
            user.new_phone = phone
            otp = await generateOTP()
            await send_sms(otp=otp, phone=phone)
            user.otp = otp
            user.save()
            print(user.otp)
            keyboard = await back_keyboard(lang)
            if lang == "uz":
                await message.answer(text=f"<b>{user.new_phone}</b> raqamiga yuborilgan tasdiqlash kodini kiriting", parse_mode='HTML', reply_markup=keyboard)
            if lang == "en":
                await message.answer(text=f"Введите код подтверждения, отправленный на номер <b>{user.new_phone}</b>.", parse_mode='HTML', reply_markup=keyboard)
            if lang == "ru":
                await message.answer(text=f"Enter the verification code sent to <b>{user.new_phone}</b>", parse_mode='HTML', reply_markup=keyboard)
            await state.set_state("get_otp_settings")
        else:
            markup = await phone_keyboard(lang)
            if lang == "uz":
                await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing👇", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_phone_number_settings")            
    else:
        lang = await get_lang(message.from_user.id)
        # if message.text == "⬅️️  Назад" or message.text == "⬅️️  Orqaga" or message.text == "⬅️️  Back":
        markup = await settings_keyboard(lang)
        if lang == "uz":
            await message.answer(text="Kerakli buyruqni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer(text="Click the required button 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer(text="Выберите нужную команду 👇", reply_markup=markup, protect_content=True)
        await state.set_state("settings")


@dp.message_handler(content_types=types.ContentTypes.TEXT, state="get_otp_settings")
async def get_phone(message: types.Message, state: FSMContext):
    user = await get_user(message.from_user.id)
    lang = user.lang
    if "⬅️️" in message.text: 
        markup = await phone_keyboard(lang)
        if lang == "uz":
            await message.answer("Telefon raqamininfizni xalqaro formatda(<b>998YYXXXXXXX</b>) kiriting. Yoki raqamni ulashing👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Enter your phone number in international format (<b>998YYXXXXXX</b>). Or share the number 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Введите свой номер телефона в международном формате (<b>998YYXXXXXX</b>). Или поделитесь номером👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_phone_number_settings")            
    else:
        if message.text == user.otp:
            user.phone = user.new_phone
            user.save()
            markup = await settings_keyboard(lang)
            if lang == "uz":
                await message.answer("✅ Telefon raqami o'zgartirildi. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer("✅Phone number has been changed. Choose the section you want👇", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer("✅ Номер телефона изменен. Пожалуйста, выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
            await state.set_state("settings")
        else:
            lang = await get_lang(message.from_user.id)
            markup = await back_keyboard(lang)
            if lang == "uz":
                await message.answer("⚠️ Yuborilgan tasdiqlash kodi xato. Qayta urinib ko'ring", reply_markup=markup, protect_content=True)
            elif lang == "en":
                await message.answer("⚠️ The verification code sent is incorrect. Try again", reply_markup=markup, protect_content=True)
            elif lang == "ru":
                await message.answer("⚠️ Присланный проверочный код неверный. Попробуйте еще раз", reply_markup=markup, protect_content=True)
            await state.set_state("get_otp_settings")

 
@dp.message_handler(state="set_lang")
async def set_language(message: types.Message, state: FSMContext):
    data = message.text
    user = await get_user(message.from_user.id)
    if message.text == "🇺🇿 O'zbek tili":
        data = "uz"
    elif message.text == "🇺🇸 English":
        data = "en"
    elif message.text == "🇷🇺 Русский язык":
        data = "ru"
    user.lang = data
    user.save()
    lang = await get_lang(message.from_user.id)
    markup = await settings_keyboard(lang)
    if lang == "uz":
        await message.answer("Til o'zgariltirildi ✅.\nKerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
    elif lang == "en":
        await message.answer("The language has been changed ✅.\nClick the required button 👇", reply_markup=markup, protect_content=True)
    elif lang == "ru":
        await message.answer("Язык изменен ✅.\nНажмите нужную кнопку👇", reply_markup=markup, protect_content=True)
    await state.set_state("settings")


@dp.message_handler(state="import_product_name", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Botimizga xush kelibsiz. Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Добро пожаловать в наш бот. Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Welcome to our bot. Please select the desired section 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    else:
        await state.update_data(import_product_name=message.text)
        back_key = await asd_back_keyboard(lang)
        if lang == "uz":
            await message.answer("Maxsulot TNV ED kodini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the product HS CODE 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите код ТН ВЭД продукта 👇", reply_markup=back_key)
        await state.set_state("import_product_acd")   
        
    
@dp.message_handler(state="import_product_acd", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        if lang == "uz":
            await message.answer("Maxsulot nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the product name 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название продукта 👇", reply_markup=back_key)
        await state.set_state("import_product_name")   
    elif message.text in ["Tashlab ketish ➡️", "Skip ➡️", "Пропустить ➡️"]:
        await state.update_data(import_product_acd="Mavjud emas")
        if lang == "uz":
            await message.answer("Import qilinayotgan davlat nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the name of the exporting country 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название страны-экспортера 👇", reply_markup=back_key)
        await state.set_state("import_country")           
    else:
        await state.update_data(import_product_acd=message.text)
        if lang == "uz":
            await message.answer("Import qilinayotgan davlat nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the name of the exporting country 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название страны-экспортера 👇", reply_markup=back_key)
        await state.set_state("import_country")   
        

@dp.message_handler(state="import_country", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        if lang == "uz":
            markup = await asd_back_keyboard(lang)
            await message.answer("Maxsulot TN VED kodini kiriting 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer("Enter the product TN VED code 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer("Введите код продукта TN VED 👇", reply_markup=markup, protect_content=True)
        await state.set_state("import_product_acd")   
    else:
        markup = await user_menu(lang)
        await state.update_data(import_country=message.text)
        if lang == "uz":
            await message.answer("Maxsulotning import narxini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the import price of the product 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите импортную цену товара 👇", reply_markup=back_key)
        await state.set_state("get_import_price")   


@dp.message_handler(state="get_import_price", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        if lang == "uz":
            await message.answer("Import qilinayotgan davlat nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the name of the exporting country 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название страны-экспортера 👇", reply_markup=back_key)
        await state.set_state("import_country")
    else:
        if message.text.isdigit():
            markup = await get_phone_keyboard(lang)
            data = await state.get_data()
            await state.update_data(import_price=message.text)
            import_product_name = data["import_product_name"]
            import_product_acd = data["import_product_acd"]
            import_country =  data["import_country"]
            import_price = message.text
            document = Document()
            user = await get_user(message.from_user.id)
            document.add_heading(f'Import uchun', 0)
            document.add_paragraph(f"Mijoz: {user.name}                 Telefon: {user.phone}")
            document.add_paragraph(f"Firma: {user.company}")
            document.add_paragraph(f"Biznes sohasi: {user.product_category.name_uz}")
            document.add_paragraph(f"Obyom: {user.monthly}")
            document.add_paragraph(f"")
            document.add_paragraph(f"Maxsulot: {import_product_name}")
            document.add_paragraph(f"ACD: {import_product_acd}")
            document.add_paragraph(f"Davlat: {import_country}")
            document.add_paragraph(f"Narxi: {import_price}")
            if lang == "uz":
                await message.answer("Ma'lumotlarqabul qilindi✅. Konsultatsiya uchun qo'ng'iroq buyurtma qilasizmi 👇", reply_markup=markup, protect_content=True)
            if lang == "en":
                await message.answer("Information received. Would you like to order a call for a consultation 👇", reply_markup=markup, protect_content=True)
            if lang == "ru":
                await message.answer("Информация получена. Хотите заказать звонок для консультации 👇", reply_markup=markup, protect_content=True)
            await state.set_state("get_import_phone")
            document.save(f'offer.docx')
            os.system("abiword --to=pdf" +str(" ") + "offer.docx")
            doc = open('./offer.pdf', 'rb')
            await bot.send_document(chat_id=-838866316, document=doc, caption=f"Imort uchun")
        else:
            if lang == "uz":
                await message.answer("Maxsulotning import narxini raqamlarda kiriting 👇", reply_markup=back_key)
            if lang == "en":
                await message.answer("Enter the import price of the product in numbers 👇", reply_markup=back_key)
            if lang == "ru":
                await message.answer("Введите импортную цену товара цифрами 👇", reply_markup=back_key)
            await state.set_state("get_import_price")   
            

@dp.message_handler(state="get_import_phone", content_types=types.ContentTypes.TEXT)
async def get_import_phone(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["Savol qoldirish", "Leave a question", "Оставить вопрос"]:
        if lang == "uz":
            await message.answer("Savolingizni qoldiring 👇", reply_markup=ReplyKeyboardRemove())
        if lang == "en":
            await message.answer("Leave your question 👇", reply_markup=ReplyKeyboardRemove())
        if lang == "ru":
            await message.answer("Оставьте свой вопрос 👇", reply_markup=ReplyKeyboardRemove())
        await state.set_state("get_contract_question")
    if message.text in ["Qo'ng'iroq buyurtma qilish", "Order a call", "Заказ звонка"]:
        markup = await user_menu(lang)
        user = await get_user(message.from_user.id)
        await bot.send_message(chat_id=-838866316, text=f"{user.name}\n\nTelefon: {user.phone}\n\n Import shartnomasi uchun")
        if lang == "uz":
            await message.answer(f"Hodimimiz siz bilan 24 soat ichida bog'lanadi", reply_markup=markup, protect_content=True)
            await message.answer(f"Bosh menyu. Kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(f"Our staff will contact you within 24 hours", reply_markup=markup, protect_content=True)
            await message.answer(f"Main menu. Select the desired section 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(f"Наши сотрудники свяжутся с вами в течение 24 часов", reply_markup=markup, protect_content=True)
            await message.answer(f"Главное меню. Выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    

@dp.message_handler(state="get_export_phone", content_types=types.ContentTypes.TEXT)
async def get_export_phone(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["Savol qoldirish", "Leave a question", "Оставить вопрос"]:
        if lang == "uz":
            await message.answer("Savolingizni qoldiring 👇", reply_markup=ReplyKeyboardRemove())
        if lang == "en":
            await message.answer("Leave your question 👇", reply_markup=ReplyKeyboardRemove())
        if lang == "ru":
            await message.answer("Оставьте свой вопрос 👇", reply_markup=ReplyKeyboardRemove())
        await state.set_state("get_contract_question")
    if message.text in ["Qo'ng'iroq buyurtma qilish", "Order a call", "Заказ звонка"]:
        markup = await user_menu(lang)
        user = await get_user(message.from_user.id)
        await bot.send_message(chat_id=-838866316, text=f"Export shartnomasi\n\n{user.name}\n\nTelefon: {user.phone}")
        if lang == "uz":
            await message.answer(f"Hodimimiz siz bilan 24 soat ichida bog'lanadi", reply_markup=markup, protect_content=True)
            await message.answer(f"Bosh menyu. Kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer(f"Our staff will contact you within 24 hours", reply_markup=markup, protect_content=True)
            await message.answer(f"Main menu. Select the desired section 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer(f"Наши сотрудники свяжутся с вами в течение 24 часов", reply_markup=markup, protect_content=True)
            await message.answer(f"Главное меню. Выберите нужный раздел 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")

            
@dp.message_handler(state="export_product_name", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        markup = await user_menu(lang)
        if lang == "uz":
            await message.answer("Iltimos kerakli bo'limni tanlang 👇", reply_markup=markup, protect_content=True)
        elif lang == "ru":
            await message.answer("Выберите нужный раздел👇", reply_markup=markup, protect_content=True)
        elif lang == "en":
            await message.answer("Please select the desired section 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_category")
    else:
        await state.update_data(export_product_name=message.text)
        back_key = await back_keyboard(lang)
        if lang == "uz":
            await message.answer("Maxsulot TN VED kodini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the product TN VED code 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите код продукта TN VED 👇", reply_markup=back_key)
        await state.set_state("export_product_acd")   


@dp.message_handler(state="export_product_acd", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        if lang == "uz":
            await message.answer("Maxsulot nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the product name 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название продукта 👇", reply_markup=back_key)
        await state.set_state("export_product_name")   
    else:
        await state.update_data(export_product_acd=message.text)
        if lang == "uz":
            await message.answer("Export qilinayotgan davlat nomini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the name of the country you are exporting to 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите название страны, в которую вы экспортируете 👇", reply_markup=back_key)
        await state.set_state("export_country")   


@dp.message_handler(state="export_country", content_types=types.ContentTypes.TEXT)
async def get_service_category(message: types.Message, state: FSMContext):
    lang = await get_lang(message.from_user.id)
    back_key = await back_keyboard(lang)
    if message.text in ["⬅️ Orqaga", "⬅️ Back", "⬅️ Назад"]:
        if lang == "uz":
            await message.answer("Maxsulot TN VED kodini kiriting 👇", reply_markup=back_key)
        if lang == "en":
            await message.answer("Enter the product TN VED code 👇", reply_markup=back_key)
        if lang == "ru":
            await message.answer("Введите код продукта TN VED 👇", reply_markup=back_key)
        await state.set_state("export_product_acd")
    else:
        markup = await user_menu(lang)
        data = await state.get_data()
        await state.update_data(import_price=message.text)
        export_product_name = data["export_product_name"]
        export_product_acd = data["export_product_acd"]
        export_country = message.text
        document = Document()
        user = await get_user(message.from_user.id)
        document.add_heading(f'Import uchun', 0)
        document.add_paragraph(f"Mijoz: {user.name}                 Telefon: {user.phone}")
        document.add_paragraph(f"Firma: {user.company}")
        document.add_paragraph(f"Biznes sohasi: {user.product_category.name_uz}")
        document.add_paragraph(f"Obyom: {user.monthly}")
        document.add_paragraph(f"")
        document.add_paragraph(f"Maxsulot: {export_product_name}")
        document.add_paragraph(f"ACD: {export_product_acd}")
        document.add_paragraph(f"Davlat: {export_country}")
        if lang == "uz":
            await message.answer("Ma'lumotlarqabul qilindi ✅. Konsultatsiya uchun qo'ng'iroq buyurtma qilasizmi 👇", reply_markup=markup, protect_content=True)
        if lang == "en":
            await message.answer("Information received ✅. Would you like to order a call for a consultation 👇", reply_markup=markup, protect_content=True)
        if lang == "ru":
            await message.answer("Информация получена ✅. Хотите заказать звонок для консультации 👇", reply_markup=markup, protect_content=True)
        await state.set_state("get_export_phone")
        await state.set_state("get_category")
        document.save(f'offer.docx')
        os.system("abiword --to=pdf" +str(" ") + "offer.docx")
        doc = open('./offer.pdf', 'rb')
        await bot.send_document(chat_id=-838866316, document=doc, caption=f"Export uchun")
           

